import io

from docx import Document as DocxDocument
from pypdf import PdfWriter

from app.file_upload.text_extract import extract_document_text


def test_extracts_markdown_as_text(tmp_path):
    p = tmp_path / "adr.md"
    p.write_text("# Decision\n\nWe use asyncpg, not psycopg2.", encoding="utf-8")
    assert "asyncpg" in extract_document_text(str(p), "adr.md")


def test_extracts_plain_text(tmp_path):
    p = tmp_path / "notes.txt"
    p.write_text("auth is validated in middleware", encoding="utf-8")
    assert "middleware" in extract_document_text(str(p), "notes.txt")


def test_extracts_docx_paragraphs(tmp_path):
    p = tmp_path / "spec.docx"
    doc = DocxDocument()
    doc.add_paragraph("Tokens are validated in middleware, never per-route.")
    doc.add_paragraph("We use asyncpg.")
    doc.save(str(p))

    text = extract_document_text(str(p), "spec.docx")
    assert "middleware" in text and "asyncpg" in text


def test_extension_wins_over_path(tmp_path):
    # The dispatch keys on the original filename's suffix, not the temp path
    # (which mkstemp gives a random suffix). A .docx spooled to a suffixless temp
    # file must still be read as docx.
    doc = DocxDocument()
    doc.add_paragraph("decision recorded here")
    tmp = tmp_path / "spool-xyz"
    doc.save(str(tmp))
    assert "decision recorded" in extract_document_text(str(tmp), "real.docx")


def test_unreadable_pdf_degrades_to_empty(tmp_path):
    # A structurally-valid but text-less PDF (no text layer) yields "".
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    p = tmp_path / "scan.pdf"
    p.write_bytes(buf.getvalue())
    assert extract_document_text(str(p), "scan.pdf") == ""


def test_binary_junk_is_empty_ish(tmp_path):
    # A .pdf that isn't a PDF at all must not raise — best-effort returns "".
    p = tmp_path / "broken.pdf"
    p.write_bytes(b"\x00\x01not a pdf\xff")
    assert extract_document_text(str(p), "broken.pdf") == ""
